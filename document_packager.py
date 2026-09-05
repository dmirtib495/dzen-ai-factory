from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


MIN_APPROVED_IMAGES = 3
MAX_APPROVED_IMAGES = 5


def _safe_name(value: str, fallback: str = "article") -> str:
    value = re.sub(r"[^\w\-а-яА-ЯёЁ ]+", "", (value or "").strip(), flags=re.UNICODE)
    value = re.sub(r"\s+", "_", value).strip("_")
    return value[:80] or fallback


def _validate_image_count(images: list[Path]) -> None:
    if len(images) < MIN_APPROVED_IMAGES or len(images) > MAX_APPROVED_IMAGES:
        raise ValueError(
            f"Approved image count must be {MIN_APPROVED_IMAGES}-{MAX_APPROVED_IMAGES}; got {len(images)}"
        )


def _add_markdown_line(document: Document, raw: str) -> None:
    line = (raw or "").strip()
    if not line:
        document.add_paragraph()
        return
    if line.startswith("## "):
        document.add_heading(line[3:].strip(), level=2)
        return
    if line.startswith("### "):
        document.add_heading(line[4:].strip(), level=3)
        return
    if re.match(r"^\d+\.\s+", line):
        p = document.add_paragraph(style="List Number")
        p.add_run(re.sub(r"^\d+\.\s+", "", line))
        return
    if line.startswith("- "):
        p = document.add_paragraph(style="List Bullet")
        p.add_run(line[2:].strip())
        return
    document.add_paragraph(line)


def _content_lines(markdown: str) -> list[str]:
    out: list[str] = []
    blank = False
    for raw in (markdown or "").splitlines():
        line = raw.strip()
        if not line:
            if out and not blank:
                out.append("")
            blank = True
            continue
        out.append(line)
        blank = False
    while out and not out[-1]:
        out.pop()
    return out


def _paragraph_indices(lines: list[str]) -> list[int]:
    return [
        i
        for i, line in enumerate(lines)
        if line and not line.startswith(("## ", "### ", "- ")) and not re.match(r"^\d+\.\s+", line)
    ]


def _nearest_paragraph_at_or_after(paragraphs: list[int], target: int, used: set[int]) -> int:
    for idx in paragraphs:
        if idx >= target and idx not in used:
            return idx
    for idx in reversed(paragraphs):
        if idx not in used:
            return idx
    return target


def _last_content_before(lines: list[str], heading_index: int) -> int | None:
    for idx in range(heading_index - 1, -1, -1):
        if lines[idx].strip():
            return idx
    return None


def _heading_index(lines: list[str], keywords: tuple[str, ...]) -> int | None:
    for i, line in enumerate(lines):
        low = line.lower()
        if line.startswith(("## ", "### ")) and any(k in low for k in keywords):
            return i
    return None


def _image_anchor_indices(lines: list[str], count: int) -> list[int]:
    """Choose editorial anchors while preserving image order around checklists/conclusions."""
    paragraphs = _paragraph_indices(lines)
    if not paragraphs:
        return [max(0, len(lines) - 1)] * count

    last = paragraphs[-1]
    ratios_by_count = {
        3: (0.05, 0.50, 0.82),
        4: (0.05, 0.34, 0.62, 0.84),
        5: (0.05, 0.31, 0.52, 0.70, 0.88),
    }
    anchors: list[int] = []
    used: set[int] = set()
    for pos, ratio in enumerate(ratios_by_count[count]):
        target = paragraphs[0] if pos == 0 else round(last * ratio)
        anchor = _nearest_paragraph_at_or_after(paragraphs, target, used)
        used.add(anchor)
        anchors.append(anchor)

    if count >= 4:
        interior = _heading_index(
            lines,
            ("салон", "интерьер", "оснащ", "комплектац", "эргоном", "мультимед", "оборудован"),
        )
        if interior is not None:
            anchors[3] = _nearest_paragraph_at_or_after(paragraphs, interior + 1, set(anchors[:3]))

    if count >= 4:
        checklist = _heading_index(lines, ("чек-лист", "чеклист", "памятк"))
        if checklist is not None:
            before_checklist = _last_content_before(lines, checklist)
            if before_checklist is not None and before_checklist > anchors[2]:
                anchors[3] = before_checklist

    if count >= 5:
        conclusion = _heading_index(lines, ("итог", "вывод", "вердикт", "заключ"))
        if conclusion is not None:
            before_conclusion = _last_content_before(lines, conclusion)
            if before_conclusion is not None and before_conclusion > anchors[3]:
                anchors[4] = before_conclusion

    normalized: list[int] = []
    floor = -1
    for anchor in anchors:
        chosen = anchor
        if chosen <= floor:
            later = [p for p in paragraphs if p > floor]
            chosen = later[0] if later else floor + 1
        normalized.append(chosen)
        floor = chosen
    return normalized


def _add_image(document: Document, image: Path, idx: int, caption: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image), width=Inches(4.55))

    text = (caption or "").strip()
    if text:
        cp = document.add_paragraph(f"Рис. {idx}. {text}")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(7)
        for run in cp.runs:
            run.italic = True
            run.font.size = Pt(9)


def build_article_docx(
    *,
    headline: str,
    article_markdown: str,
    approved_images: Iterable[str | Path],
    output_path: str | Path,
    captions: Iterable[str] | None = None,
) -> Path:
    images = [Path(p) for p in approved_images]
    _validate_image_count(images)
    missing = [str(p) for p in images if not p.is_file()]
    if missing:
        raise FileNotFoundError("Approved image files missing: " + ", ".join(missing))

    caption_list = list(captions or [])
    while len(caption_list) < len(images):
        caption_list.append("")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run((headline or "").strip())
    run.bold = True
    run.font.size = Pt(18)

    lines = _content_lines(article_markdown)
    anchors = _image_anchor_indices(lines, len(images))
    images_after: dict[int, list[int]] = {}
    for image_idx, anchor in enumerate(anchors):
        images_after.setdefault(anchor, []).append(image_idx)

    for line_idx, line in enumerate(lines):
        _add_markdown_line(doc, line)
        for image_idx in images_after.get(line_idx, []):
            _add_image(doc, images[image_idx], image_idx + 1, caption_list[image_idx])

    inserted = {i for indexes in images_after.values() for i in indexes}
    for image_idx, image in enumerate(images):
        if image_idx not in inserted:
            _add_image(doc, image, image_idx + 1, caption_list[image_idx])

    doc.save(output)
    if not output.is_file() or output.stat().st_size <= 0:
        raise RuntimeError("DOCX was not created correctly")
    return output


def build_article_package(
    *,
    article_id: int | str,
    headline: str,
    article_markdown: str,
    approved_images: Iterable[str | Path],
    output_root: str | Path = "data/packages",
    captions: Iterable[str] | None = None,
) -> tuple[Path, Path, Path]:
    images = [Path(p) for p in approved_images]
    _validate_image_count(images)

    root = Path(output_root)
    folder = root / f"article_{article_id}_{_safe_name(headline)}"
    if folder.exists():
        shutil.rmtree(folder)
    image_dir = folder / "images"
    image_dir.mkdir(parents=True, exist_ok=True)

    copied: list[Path] = []
    for idx, src in enumerate(images, 1):
        if not src.is_file():
            raise FileNotFoundError(str(src))
        suffix = src.suffix.lower() if src.suffix else ".jpg"
        dst = image_dir / f"image_{idx:02d}{suffix}"
        shutil.copy2(src, dst)
        copied.append(dst)

    docx_path = folder / "article.docx"
    build_article_docx(
        headline=headline,
        article_markdown=article_markdown,
        approved_images=copied,
        output_path=docx_path,
        captions=captions,
    )

    (folder / "article.md").write_text(article_markdown or "", encoding="utf-8")
    (folder / "headline.txt").write_text((headline or "").strip(), encoding="utf-8")
    (folder / "package_manifest.json").write_text(
        json.dumps({"article_id": int(article_id), "image_count": len(copied)}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    zip_path = root / f"{folder.name}.zip"
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(folder.rglob("*")):
            if path.is_file():
                zf.write(path, arcname=str(path.relative_to(folder)))

    if not zip_path.is_file() or zip_path.stat().st_size <= 0:
        raise RuntimeError("ZIP was not created correctly")
    return folder, docx_path, zip_path
